"""DOCX reflow renderer — editable document preserving semantic structure.

Figures are embedded as **inline crops** at their reading_order position
(per resolved design decision — no sidecar).
"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Inches
from PIL import Image

from ocr_agentic_engine.renderers._ordering import iter_blocks
from ocr_agentic_engine.types import Annotations, DocumentRepresentation, OCRBlock


def render_docx_reflow(doc: DocumentRepresentation, ann: Annotations) -> bytes:
    d = Document()
    if ann.title is not None and 0 <= ann.title < len(doc.blocks):
        d.add_heading(ann.corrections.get(ann.title, doc.blocks[ann.title].text), level=1)

    current_list_id: int | None = None
    for idx, b in iter_blocks(doc, ann, redact_handwritten=False):
        if idx == ann.title:
            continue
        if idx in ann.handwritten:
            d.add_paragraph("[handwritten content redacted]").italic = True
            continue
        text = ann.corrections.get(idx, b.text)

        if idx in ann.headings:
            d.add_heading(text, level=min(3, ann.headings[idx] + 1))
        elif idx in ann.lists:
            d.add_paragraph(text, style="List Bullet")
            current_list_id = ann.lists[idx]
        elif b.kind == "table" and b.table_cells:
            _write_table(d, b)
        elif b.kind == "figure":
            _inline_figure(d, doc, b)
        elif b.kind in ("header", "footer"):
            p = d.add_paragraph(text); p.runs[0].italic = True
        else:
            d.add_paragraph(text)
            current_list_id = None

    buf = BytesIO(); d.save(buf); return buf.getvalue()


def _write_table(d: Document, b: OCRBlock) -> None:
    rows: dict[int, dict[int, str]] = {}
    for c in b.table_cells or []:
        rows.setdefault(c.row, {})[c.col] = c.text
    if not rows:
        d.add_paragraph(b.text); return
    n_rows = max(rows) + 1
    n_cols = max((max(r) for r in rows.values()), default=0) + 1
    tbl = d.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    for r, cols in rows.items():
        for c, text in cols.items():
            tbl.cell(r, c).text = text


def _inline_figure(d: Document, doc: DocumentRepresentation, b: OCRBlock) -> None:
    pi = next((p for p in doc.page_images if p.page == b.page), None)
    if pi is None:
        d.add_paragraph("[figure]"); return
    try:
        with Image.open(pi.path) as im:
            x0, y0, x1, y1 = b.bbox
            crop = im.crop((max(0, x0), max(0, y0), min(im.width, x1), min(im.height, y1)))
            buf = BytesIO(); crop.save(buf, format="PNG"); buf.seek(0)
        d.add_picture(buf, width=Inches(4.5))
    except Exception:
        d.add_paragraph("[figure]")
